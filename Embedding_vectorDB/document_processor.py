# Embedding_vectorDB/document_processor.py - ENHANCED WITH SMART CHUNKING

import PyPDF2
from docx import Document
import pandas as pd
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import re
from typing import List, Dict
import logging

# Import Docling processor
from docling_processor import get_docling_processor

logger = logging.getLogger(__name__)


# ================================================================
# SMART CHUNKING STRATEGY
# ================================================================

class SmartChunker:
    """
    Chiến lược chunking thông minh cho văn bản tiếng Việt

    Ưu tiên:
    1. Giữ nguyên ngữ cảnh section/heading
    2. Chia theo đoạn văn logic
    3. Đảm bảo chunk size hợp lý (300-800 tokens)
    4. Overlap giữa các chunks để giữ ngữ cảnh
    """

    def __init__(
            self,
            target_chunk_size: int = 500,
            min_chunk_size: int = 200,
            max_chunk_size: int = 800,
            overlap_size: int = 100
    ):
        self.target_chunk_size = target_chunk_size
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.overlap_size = overlap_size

    def estimate_tokens(self, text: str) -> int:
        """Ước lượng số tokens (Vietnamese words)"""
        words = re.findall(r'\w+|[^\w\s]', text)
        return len(words)

    def parse_markdown_structure(self, markdown_content: str) -> List[Dict]:
        """Parse markdown thành cấu trúc phân cấp"""
        lines = markdown_content.split('\n')
        sections = []
        current_section = {
            'level': 0,
            'title': 'Document Root',
            'content': [],
            'parent_titles': []
        }

        for line in lines:
            line_stripped = line.strip()

            if line_stripped.startswith('#'):
                if current_section['content']:
                    sections.append(current_section.copy())

                heading_match = re.match(r'^(#{1,6})\s+(.+)$', line_stripped)
                if heading_match:
                    level = len(heading_match.group(1))
                    title = heading_match.group(2).strip()

                    parent_titles = []
                    for sec in reversed(sections):
                        if sec['level'] < level:
                            parent_titles.insert(0, sec['title'])
                            if sec['parent_titles']:
                                parent_titles = sec['parent_titles'] + parent_titles
                            break

                    current_section = {
                        'level': level,
                        'title': title,
                        'content': [],
                        'parent_titles': parent_titles
                    }

            elif line_stripped:
                current_section['content'].append(line_stripped)

        if current_section['content']:
            sections.append(current_section)

        return sections

    def create_semantic_chunks(self, markdown_content: str) -> List[Dict]:
        """Tạo chunks dựa trên semantic structure"""
        sections = self.parse_markdown_structure(markdown_content)
        chunks = []

        for section in sections:
            section_text = '\n'.join(section['content'])
            section_tokens = self.estimate_tokens(section_text)

            context_path = section['parent_titles'] + [section['title']]
            context_header = '\n'.join([
                f"{'#' * (i + 1)} {title}"
                for i, title in enumerate(context_path)
            ])

            # Case 1: Section nhỏ - giữ nguyên
            if section_tokens <= self.max_chunk_size:
                chunks.append({
                    'section_title': ' > '.join(context_path),
                    'content': f"{context_header}\n\n{section_text}",
                    'context': context_header,
                    'context_path': context_path,
                    'level': section['level'],
                    'token_count': section_tokens,
                    'chunk_type': 'complete_section'
                })

            # Case 2: Section lớn - chia nhỏ
            else:
                sub_chunks = self._split_large_section(
                    section_text,
                    context_header,
                    section['title'],
                    context_path,
                    section['level']
                )
                chunks.extend(sub_chunks)

        # Add overlap between chunks
        chunks = self._add_overlap(chunks)

        logger.info(f"✅ Created {len(chunks)} semantic chunks")
        return chunks

    def _split_large_section(
            self,
            text: str,
            context_header: str,
            section_title: str,
            context_path: List[str],
            level: int
    ) -> List[Dict]:
        """Chia section lớn thành chunks nhỏ hơn"""
        chunks = []
        paragraphs = re.split(r'\n\n+', text)

        current_chunk = []
        current_tokens = 0
        chunk_index = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_tokens = self.estimate_tokens(para)

            # Đoạn văn quá lớn - chia theo câu
            if para_tokens > self.max_chunk_size:
                if current_chunk:
                    chunk_text = '\n\n'.join(current_chunk)
                    chunks.append({
                        'section_title': ' > '.join(context_path) + f" (Part {chunk_index + 1})",
                        'content': f"{context_header}\n\n{chunk_text}",
                        'context': context_header,
                        'context_path': context_path,
                        'level': level,
                        'token_count': current_tokens,
                        'chunk_type': 'partial_section',
                        'part_index': chunk_index
                    })
                    current_chunk = []
                    current_tokens = 0
                    chunk_index += 1

                sub_chunks = self._split_long_paragraph(
                    para, context_header, section_title,
                    context_path, level, chunk_index
                )
                chunks.extend(sub_chunks)
                chunk_index += len(sub_chunks)

            # Thêm vào chunk hiện tại
            elif current_tokens + para_tokens <= self.target_chunk_size:
                current_chunk.append(para)
                current_tokens += para_tokens

            # Chunk đầy - tạo mới
            else:
                if current_chunk:
                    chunk_text = '\n\n'.join(current_chunk)
                    chunks.append({
                        'section_title': ' > '.join(context_path) + f" (Part {chunk_index + 1})",
                        'content': f"{context_header}\n\n{chunk_text}",
                        'context': context_header,
                        'context_path': context_path,
                        'level': level,
                        'token_count': current_tokens,
                        'chunk_type': 'partial_section',
                        'part_index': chunk_index
                    })
                    chunk_index += 1

                current_chunk = [para]
                current_tokens = para_tokens

        # Chunk cuối
        if current_chunk:
            chunk_text = '\n\n'.join(current_chunk)
            chunks.append({
                'section_title': ' > '.join(context_path) + f" (Part {chunk_index + 1})",
                'content': f"{context_header}\n\n{chunk_text}",
                'context': context_header,
                'context_path': context_path,
                'level': level,
                'token_count': current_tokens,
                'chunk_type': 'partial_section',
                'part_index': chunk_index
            })

        return chunks

    def _split_long_paragraph(
            self, paragraph: str, context_header: str,
            section_title: str, context_path: List[str],
            level: int, start_index: int
    ) -> List[Dict]:
        """Chia đoạn văn dài theo câu"""
        chunks = []
        sentences = re.split(r'([.!?]+\s+)', paragraph)

        full_sentences = []
        i = 0
        while i < len(sentences):
            if i + 1 < len(sentences) and re.match(r'[.!?]+\s+', sentences[i + 1]):
                full_sentences.append(sentences[i] + sentences[i + 1])
                i += 2
            else:
                if sentences[i].strip():
                    full_sentences.append(sentences[i])
                i += 1

        current_chunk = []
        current_tokens = 0
        chunk_index = start_index

        for sentence in full_sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            sent_tokens = self.estimate_tokens(sentence)

            if current_tokens + sent_tokens <= self.target_chunk_size:
                current_chunk.append(sentence)
                current_tokens += sent_tokens
            else:
                if current_chunk:
                    chunk_text = ' '.join(current_chunk)
                    chunks.append({
                        'section_title': ' > '.join(context_path) + f" (Part {chunk_index + 1})",
                        'content': f"{context_header}\n\n{chunk_text}",
                        'context': context_header,
                        'context_path': context_path,
                        'level': level,
                        'token_count': current_tokens,
                        'chunk_type': 'sentence_group',
                        'part_index': chunk_index
                    })
                    chunk_index += 1

                current_chunk = [sentence]
                current_tokens = sent_tokens

        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                'section_title': ' > '.join(context_path) + f" (Part {chunk_index + 1})",
                'content': f"{context_header}\n\n{chunk_text}",
                'context': context_header,
                'context_path': context_path,
                'level': level,
                'token_count': current_tokens,
                'chunk_type': 'sentence_group',
                'part_index': chunk_index
            })

        return chunks

    def _add_overlap(self, chunks: List[Dict]) -> List[Dict]:
        """Thêm overlap giữa các chunks liên tiếp"""
        if len(chunks) <= 1:
            return chunks

        overlapped_chunks = []

        for i, chunk in enumerate(chunks):
            chunk_copy = chunk.copy()

            if i > 0:
                prev_chunk = chunks[i - 1]
                if prev_chunk['context_path'] == chunk['context_path']:
                    overlap_text = self._extract_last_tokens(
                        prev_chunk['content'],
                        self.overlap_size
                    )

                    if overlap_text:
                        chunk_copy['content'] = f"{overlap_text}\n\n---\n\n{chunk_copy['content']}"

            overlapped_chunks.append(chunk_copy)

        return overlapped_chunks

    def _extract_last_tokens(self, text: str, num_tokens: int) -> str:
        """Trích xuất N tokens cuối từ text"""
        sentences = re.split(r'([.!?]+)', text)

        result = []
        current_tokens = 0

        for i in range(len(sentences) - 1, -1, -1):
            sent = sentences[i].strip()
            if not sent:
                continue

            sent_tokens = self.estimate_tokens(sent)

            if current_tokens + sent_tokens <= num_tokens:
                result.insert(0, sent)
                current_tokens += sent_tokens
            else:
                break

        return ' '.join(result)


# ================================================================
# ENHANCED DOCUMENT PROCESSOR
# ================================================================

class DocumentProcessor:
    """
    Document processor with Docling integration and Smart Chunking
    """

    def __init__(self, use_docling: bool = True, use_ocr: bool = True):
        """
        Initialize document processor

        Args:
            use_docling: Use Docling for processing (recommended)
            use_ocr: Enable OCR for scanned documents
        """
        self.use_docling = use_docling
        self.use_ocr = use_ocr

        # Initialize Docling processor
        self.docling_processor = None
        if self.use_docling:
            try:
                self.docling_processor = get_docling_processor(use_ocr=use_ocr)
                logger.info("✅ Docling processor ready")
            except Exception as e:
                logger.warning(f"⚠️ Docling initialization failed: {e}")
                self.docling_processor = None

        # Initialize Smart Chunker
        self.chunker = SmartChunker(
            target_chunk_size=500,
            min_chunk_size=200,
            max_chunk_size=800,
            overlap_size=100
        )
        logger.info("✅ Smart Chunker initialized")

    def process_pdf(self, file_path: str) -> str:
        """Process PDF file - Docling first, then fallback"""

        # Try Docling
        if self.docling_processor:
            try:
                logger.info("🚀 Trying Docling for PDF processing...")
                markdown_content = self.docling_processor.process_pdf(file_path)

                if markdown_content and len(markdown_content.strip()) > 50:
                    logger.info("✅ Docling PDF processing successful")
                    return self.clean_and_structure_markdown(markdown_content)
            except Exception as e:
                logger.warning(f"⚠️ Docling failed: {e}, using fallback")

        # Fallback to legacy
        logger.info("📋 Using legacy PDF processing...")
        markdown_content = ""

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""

                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        page_text = self.clean_page_artifacts(page_text)
                        full_text += f" {page_text}"

                markdown_content = full_text.strip()

            if len(markdown_content.strip()) < 100:
                logger.info("🔍 Text extraction insufficient, using OCR...")
                markdown_content = self.ocr_pdf_continuous(file_path)

        except Exception as e:
            logger.error(f"❌ Legacy PDF processing failed: {e}")
            markdown_content = self.ocr_pdf_continuous(file_path)

        return self.clean_and_structure_markdown(markdown_content)

    def process_word(self, file_path: str) -> str:
        """Process Word document - Docling first, then fallback"""

        # Try Docling
        if self.docling_processor:
            try:
                logger.info("🚀 Trying Docling for DOCX processing...")
                markdown_content = self.docling_processor.process_docx(file_path)

                if markdown_content and len(markdown_content.strip()) > 50:
                    logger.info("✅ Docling DOCX processing successful")
                    return self.clean_and_structure_markdown(markdown_content)

            except Exception as e:
                logger.warning(f"⚠️ Docling failed: {e}, using fallback")

        # Fallback
        logger.info("📋 Using legacy DOCX processing...")

        try:
            doc = Document(file_path)
            markdown_content = ""

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[
                            -1].isdigit() else 1
                        markdown_content += f"\n{'#' * level} {text}\n\n"
                    else:
                        markdown_content += f"{text} "

            for table in doc.tables:
                markdown_content += "\n\n" + self.convert_table_to_markdown(table) + "\n\n"

            return self.clean_and_structure_markdown(markdown_content)

        except Exception as e:
            logger.error(f"❌ DOCX processing error: {e}")
            raise Exception(f"Error processing Word document: {e}")

    def process_excel(self, file_path: str) -> str:
        """Process Excel file"""
        try:
            logger.info("📊 Processing Excel file...")

            excel_file = pd.ExcelFile(file_path)
            markdown_content = ""

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                if not df.empty:
                    markdown_content += f"\n## {sheet_name}\n\n"
                    markdown_content += self.convert_dataframe_to_1d_markdown(df) + "\n"

            return self.clean_and_structure_markdown(markdown_content)

        except Exception as e:
            logger.error(f"❌ Excel processing error: {e}")
            raise Exception(f"Error processing Excel file: {e}")

    def process_text(self, text_content: str) -> str:
        """Process plain text"""
        cleaned_text = self.clean_page_artifacts(text_content)
        return self.clean_and_structure_markdown(cleaned_text)

    # ================================================================
    # SMART CHUNKING METHOD (REPLACES parse_markdown_to_sentences)
    # ================================================================

    def parse_markdown_to_chunks(self, markdown_content: str) -> List[Dict]:
        """
        Parse markdown thành chunks thông minh

        THAY THẾ: parse_markdown_to_sentences()
        """
        logger.info("🧠 Using Smart Chunking strategy...")
        return self.chunker.create_semantic_chunks(markdown_content)

    # ================================================================
    # LEGACY HELPER METHODS
    # ================================================================

    def ocr_pdf_continuous(self, file_path: str) -> str:
        """OCR PDF pages"""
        full_text = ""
        try:
            pages = convert_from_path(file_path)
            for page in pages:
                page_text = pytesseract.image_to_string(page, lang='vie')
                cleaned_text = self.clean_page_artifacts(page_text)
                if cleaned_text.strip():
                    full_text += f" {cleaned_text}"
        except Exception as e:
            logger.error(f"OCR Error: {e}")
        return full_text.strip()

    def clean_page_artifacts(self, text: str) -> str:
        """Remove page numbers, headers, footers"""
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                continue
            if re.match(r'^\d+$', line):
                continue
            if re.match(r'^Trang \d+', line, re.IGNORECASE):
                continue
            if re.match(r'^Page \d+', line, re.IGNORECASE):
                continue
            if len(line) < 5 and line.isdigit():
                continue
            cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    def convert_dataframe_to_1d_markdown(self, df: pd.DataFrame) -> str:
        """Convert DataFrame to markdown"""
        markdown = ""
        if df.empty:
            return markdown

        columns = df.columns.tolist()
        for index, row in df.iterrows():
            row_text = []
            for col in columns:
                value = str(row[col]) if pd.notna(row[col]) else ""
                if value.strip():
                    row_text.append(f"**{col}**: {value}")
            if row_text:
                markdown += ", ".join(row_text) + "\n\n"

        return markdown

    def convert_table_to_markdown(self, table) -> str:
        """Convert Word table to markdown"""
        markdown = ""
        headers = []
        if table.rows:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip())

        for row in table.rows[1:]:
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(headers):
                    value = cell.text.strip()
                    if value:
                        row_data.append(f"**{headers[col_idx]}**: {value}")
            if row_data:
                markdown += ", ".join(row_data) + "\n\n"

        return markdown

    def clean_and_structure_markdown(self, content: str) -> str:
        """Clean and structure markdown"""
        content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)
        content = re.sub(r'[ \t]+', ' ', content)
        content = re.sub(r'\n\d+\n', '\n', content)
        content = re.sub(r'^Trang \d+.*\n', '', content, flags=re.MULTILINE | re.IGNORECASE)
        content = re.sub(r'^Page \d+.*\n', '', content, flags=re.MULTILINE | re.IGNORECASE)

        lines = content.split('\n')
        structured_lines = []
        previous_was_heading = False

        for line in lines:
            line = line.strip()
            if not line:
                if not previous_was_heading:
                    structured_lines.append('')
                continue

            is_heading = False
            if (line.isupper() and len(line.split()) <= 10 and len(line.split()) > 1) or \
                    re.match(r'^\d+\.?\s+[A-ZÀ-Ỹ]', line):
                if not line.startswith('#'):
                    line = f"## {line}"
                    is_heading = True

            structured_lines.append(line)
            previous_was_heading = is_heading

        final_content = '\n'.join(structured_lines)
        final_content = re.sub(r'\n{3,}', '\n\n', final_content)

        return final_content.strip()